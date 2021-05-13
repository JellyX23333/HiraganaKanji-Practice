import docx, csv, random

WORDMAX = 30

wordSource = []
words = []
answers = ""

def fillWord(word):
    if len(word) < 6:
        word += (6 - len(word)) * ' '
    return word

with open("wordSelection.config", 'r') as config:
    reader = csv.reader(config)
    for line in reader:
        wordSource.append({"name":line[0], "proportion":line[1]})

for file in wordSource:
    fileWordMax = 30 * float(file["proportion"])
    fileWordList = []
    
    with open(file["name"], 'r') as wordFile:
        reader = csv.reader(wordFile)
        for line in reader:
            if line[0] != "":
                fileWordList.append({"hiragana":line[1], "kanji":line[0]})
                
    fileWordList = random.choices(fileWordList, k=int(round(float(30) * float(file["proportion"]))))
    for word in fileWordList:
        words.append(word)

document = docx.Document()
document.add_heading("Hiragana Test", 0)

random.shuffle(words)
count = 0
for i in range(int(round(float(WORDMAX) / float(6)))):
    sentence = ""
    for i in range(6):
        if count + i > WORDMAX - 1: break;
        sentence += str(count + i) + '.' + ' '+ fillWord(words[count + i]["hiragana"]) + '\t'
        answers += str(count + i) + '.' + ' ' + words[count + i]["kanji"]
    count += 6
    document.add_paragraph(sentence)
    document.add_paragraph(("("+" "*19+")\t")*6)
    
document.add_page_break()
document.add_paragraph(answers)

document.save("test.docx")
